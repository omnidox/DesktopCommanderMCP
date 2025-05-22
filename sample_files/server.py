"""
Claude Document MCP Server - Model Context Protocol server for Claude Desktop

Features:
- Microsoft Word file operations (create, edit, convert from txt, extract text)
- Excel file operations (create, edit, convert from csv)
- PDF file operations (create, convert from Word)

This is a headless server with no UI, designed to be used with Claude Desktop.
"""

import os
import sys
import json
import logging
from pathlib import Path
from typing import Dict, Any, List, Optional
from io import BytesIO


from mcp.server.fastmcp import FastMCP

# Document processing libraries
try:
    import docx
    from docx import Document
    from docx.shared import Pt, Inches
except ImportError:
    raise ImportError("Please install python-docx with: uv pip install python-docx")

try:
    import pandas as pd
    import openpyxl
except ImportError:
    raise ImportError("Please install pandas and openpyxl with: uv pip install pandas openpyxl")

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
except ImportError:
    raise ImportError("Please install reportlab with: uv pip install reportlab")

try:
    import docx2pdf
except ImportError:
    raise ImportError("Please install docx2pdf with: uv pip install docx2pdf")

# Set up logging
log_dir = Path(__file__).parent.parent / "logs"
log_dir.mkdir(exist_ok=True)
log_file = log_dir / "document_mcp.log"

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[
        logging.StreamHandler(),
        logging.FileHandler(log_file)
    ]
)
logger = logging.getLogger(__name__)

# Initialize the FastMCP server
server = FastMCP(
    "Document Operations", 
    description="MCP server for document operations (Word, Excel, PDF)",
    dependencies=[
        "python-docx", 
        "pandas", 
        "openpyxl", 
        "reportlab", 
        "docx2pdf",
    ]
)

mcp = server


# ---- Microsoft Word Operations ----

@server.tool()
def create_word_document(filepath: str, content: str) -> Dict[str, Any]:
    """
    Create a new Microsoft Word document with the provided content.
    
    Args:
        filepath: Path where to save the document
        content: Text content for the document
        
    Returns:
        Operation result with success status, message, and filepath
    """
    try:
        # Create a new document
        doc = Document()
        
        # Add content
        doc.add_paragraph(content)
        
        # Ensure the directory exists
        os.makedirs(os.path.dirname(os.path.abspath(filepath)), exist_ok=True)
        
        # Save the document
        doc.save(filepath)
        
        logger.info(f"Created Word document: {filepath}")
        return {
            "success": True,
            "message": "Successfully created Word document",
            "filepath": filepath
        }
    except Exception as e:
        logger.error(f"Error creating Word document: {str(e)}")
        return {
            "success": False,
            "message": f"Error creating Word document: {str(e)}",
            "filepath": None
        }

@server.tool()
def edit_word_document(filepath: str, operations: List[Dict[str, Any]]) -> Dict[str, Any]:
    """
    Edit an existing Microsoft Word document using the specified operations.
    
    Args:
        filepath: Path to the Word document
        operations: List of operations to perform, where each operation is a dictionary with:
            - type: Operation type (add_paragraph, add_heading, edit_paragraph, delete_paragraph)
            - Additional parameters depending on the operation type
            
    Returns:
        Operation result with success status, message, and filepath
    """
    try:
        # Load the document
        if not os.path.exists(filepath):
            return {
                "success": False,
                "message": f"File not found: {filepath}",
                "filepath": None
            }
        
        doc = Document(filepath)
        
        # Apply operations
        for op in operations:
            op_type = op.get("type")
            
            if op_type == "add_paragraph":
                doc.add_paragraph(op.get("text", ""))
            
            elif op_type == "add_heading":
                doc.add_heading(op.get("text", ""), level=op.get("level", 1))
            
            elif op_type == "edit_paragraph":
                idx = op.get("index", 0)
                new_text = op.get("text", "")
                
                if 0 <= idx < len(doc.paragraphs):
                    doc.paragraphs[idx].text = new_text
                else:
                    logger.warning(f"Paragraph index out of range: {idx}")
            
            elif op_type == "delete_paragraph":
                idx = op.get("index", 0)
                
                if 0 <= idx < len(doc.paragraphs):
                    p = doc.paragraphs[idx]
                    p_elem = p._element
                    p_elem.getparent().remove(p_elem)
                else:
                    logger.warning(f"Paragraph index out of range: {idx}")
            
            else:
                logger.warning(f"Unknown operation type: {op_type}")
        
        # Save the document
        doc.save(filepath)
        
        logger.info(f"Edited Word document: {filepath}")
        return {
            "success": True,
            "message": "Successfully edited Word document",
            "filepath": filepath
        }
    except Exception as e:
        logger.error(f"Error editing Word document: {str(e)}")
        return {
            "success": False,
            "message": f"Error editing Word document: {str(e)}",
            "filepath": None
        }

@server.tool()
def convert_txt_to_word(source_path: str, target_path: str) -> Dict[str, Any]:
    """
    Convert a text file to a Microsoft Word document.
    
    Args:
        source_path: Path to the text file
        target_path: Path where to save the Word document
        
    Returns:
        Operation result with success status, message, and filepath
    """
    try:
        # Check if source file exists
        if not os.path.exists(source_path):
            return {
                "success": False,
                "message": f"Source file not found: {source_path}",
                "filepath": None
            }
        
        # Read the text file
        with open(source_path, 'r', encoding='utf-8') as file:
            text_content = file.read()
        
        # Create a new document
        doc = Document()
        
        # Add content as paragraphs (split by newlines)
        for paragraph in text_content.split('\n'):
            if paragraph.strip():  # Skip empty paragraphs
                doc.add_paragraph(paragraph)
        
        # Ensure the directory exists
        os.makedirs(os.path.dirname(os.path.abspath(target_path)), exist_ok=True)
        
        # Save the document
        doc.save(target_path)
        
        logger.info(f"Converted text to Word: {source_path} -> {target_path}")
        return {
            "success": True,
            "message": "Successfully converted text to Word document",
            "filepath": target_path
        }
    except Exception as e:
        logger.error(f"Error converting text to Word: {str(e)}")
        return {
            "success": False,
            "message": f"Error converting text to Word: {str(e)}",
            "filepath": None
        }


# ---- Word Text Extraction ----

@server.tool()
def extract_docx_text(filepath: str) -> Dict[str, Any]:
    try:
        if not os.path.exists(filepath):
            return {
                "success": False,
                "message": f"File not found: {filepath}",
                "content": ""
            }

        doc = Document(filepath)
        full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        return {
            "success": True,
            "message": "Successfully extracted text from .docx file",
            "content": full_text
        }
    except Exception as e:
        logger.error(f"Error reading Word file: {str(e)}")
        return {
            "success": False,
            "message": f"Error reading Word file: {str(e)}",
            "content": ""
        }



# ---- Excel Operations ----

@server.tool()
def create_excel_file(filepath: str, content: str) -> Dict[str, Any]:
    """
    Create a new Excel file with the provided content.
    
    Args:
        filepath: Path where to save the Excel file
        content: Data content, either JSON string or CSV-like string
        
    Returns:
        Operation result with success status, message, and filepath
    """
    try:
        # Parse the content as JSON data
        try:
            data = json.loads(content)
        except json.JSONDecodeError:
            # If not valid JSON, treat as CSV
            data = [line.split(',') for line in content.strip().split('\n')]
        
        # Convert to DataFrame
        df = pd.DataFrame(data)
        
        # Ensure the directory exists
        os.makedirs(os.path.dirname(os.path.abspath(filepath)), exist_ok=True)
        
        # Save to Excel
        df.to_excel(filepath, index=False)
        
        logger.info(f"Created Excel file: {filepath}")
        return {
            "success": True,
            "message": "Successfully created Excel file",
            "filepath": filepath
        }
    except Exception as e:
        logger.error(f"Error creating Excel file: {str(e)}")
        return {
            "success": False,
            "message": f"Error creating Excel file: {str(e)}",
            "filepath": None
        }

@server.tool()
def edit_excel_file(filepath: str, operations: List[Dict[str, Any]]) -> Dict[str, Any]:
    """
    Edit an existing Excel file using the specified operations.
    
    Args:
        filepath: Path to the Excel file
        operations: List of operations to perform, where each operation is a dictionary with:
            - type: Operation type (update_cell, update_range, delete_row, delete_column, add_sheet, delete_sheet)
            - Additional parameters depending on the operation type
            
    Returns:
        Operation result with success status, message, and filepath
    """
    try:
        # Check if file exists
        if not os.path.exists(filepath):
            return {
                "success": False,
                "message": f"File not found: {filepath}",
                "filepath": None
            }
        
        # Load the Excel file
        wb = openpyxl.load_workbook(filepath)
        
        # Apply operations
        for op in operations:
            op_type = op.get("type")
            sheet_name = op.get("sheet", wb.sheetnames[0])
            
            # Get the sheet, create if it doesn't exist
            if sheet_name not in wb.sheetnames:
                wb.create_sheet(sheet_name)
            
            sheet = wb[sheet_name]
            
            if op_type == "update_cell":
                row = op.get("row", 1)
                col = op.get("col", 1)
                value = op.get("value", "")
                
                sheet.cell(row=row, column=col, value=value)
            
            elif op_type == "update_range":
                start_row = op.get("start_row", 1)
                start_col = op.get("start_col", 1)
                values = op.get("values", [])
                
                for i, row_values in enumerate(values):
                    for j, value in enumerate(row_values):
                        sheet.cell(row=start_row + i, column=start_col + j, value=value)
            
            elif op_type == "delete_row":
                row = op.get("row", 1)
                sheet.delete_rows(row)
            
            elif op_type == "delete_column":
                col = op.get("col", 1)
                sheet.delete_cols(col)
            
            elif op_type == "add_sheet":
                new_sheet_name = op.get("name", "NewSheet")
                if new_sheet_name not in wb.sheetnames:
                    wb.create_sheet(new_sheet_name)
            
            elif op_type == "delete_sheet":
                if sheet_name in wb.sheetnames and len(wb.sheetnames) > 1:
                    del wb[sheet_name]
            
            else:
                logger.warning(f"Unknown operation type: {op_type}")
        
        # Save the workbook
        wb.save(filepath)
        
        logger.info(f"Edited Excel file: {filepath}")
        return {
            "success": True,
            "message": "Successfully edited Excel file",
            "filepath": filepath
        }
    except Exception as e:
        logger.error(f"Error editing Excel file: {str(e)}")
        return {
            "success": False,
            "message": f"Error editing Excel file: {str(e)}",
            "filepath": None
        }

@server.tool()
def convert_csv_to_excel(source_path: str, target_path: str) -> Dict[str, Any]:
    """
    Convert a CSV file to an Excel file.
    
    Args:
        source_path: Path to the CSV file
        target_path: Path where to save the Excel file
        
    Returns:
        Operation result with success status, message, and filepath
    """
    try:
        # Check if source file exists
        if not os.path.exists(source_path):
            return {
                "success": False,
                "message": f"Source file not found: {source_path}",
                "filepath": None
            }
        
        # Read the CSV file
        df = pd.read_csv(source_path)
        
        # Ensure the directory exists
        os.makedirs(os.path.dirname(os.path.abspath(target_path)), exist_ok=True)
        
        # Save to Excel
        df.to_excel(target_path, index=False)
        
        logger.info(f"Converted CSV to Excel: {source_path} -> {target_path}")
        return {
            "success": True,
            "message": "Successfully converted CSV to Excel",
            "filepath": target_path
        }
    except Exception as e:
        logger.error(f"Error converting CSV to Excel: {str(e)}")
        return {
            "success": False,
            "message": f"Error converting CSV to Excel: {str(e)}",
            "filepath": None
        }

# ---- PDF Operations ----

@server.tool()
def create_pdf_file(filepath: str, content: str) -> Dict[str, Any]:
    """
    Create a new PDF file with the provided text content.
    
    Args:
        filepath: Path where to save the PDF file
        content: Text content for the PDF
        
    Returns:
        Operation result with success status, message, and filepath
    """
    try:
        # Ensure the directory exists
        os.makedirs(os.path.dirname(os.path.abspath(filepath)), exist_ok=True)
        
        # Create a new PDF with ReportLab
        c = canvas.Canvas(filepath, pagesize=letter)
        width, height = letter
        
        # Process text content
        lines = content.split('\n')
        
        y_position = height - 40  # Start position from top
        for line in lines:
            if y_position < 40:  # If we're at the bottom of the page
                c.showPage()  # Create a new page
                y_position = height - 40  # Reset position
            
            c.drawString(40, y_position, line)
            y_position -= 15  # Move down for next line
        
        c.save()
        
        logger.info(f"Created PDF file: {filepath}")
        return {
            "success": True,
            "message": "Successfully created PDF file",
            "filepath": filepath
        }
    except Exception as e:
        logger.error(f"Error creating PDF file: {str(e)}")
        return {
            "success": False,
            "message": f"Error creating PDF file: {str(e)}",
            "filepath": None
        }

@server.tool()
def convert_word_to_pdf(source_path: str, target_path: str) -> Dict[str, Any]:
    """
    Convert a Microsoft Word document to a PDF file.
    
    Args:
        source_path: Path to the Word document
        target_path: Path where to save the PDF file
        
    Returns:
        Operation result with success status, message, and filepath
    """
    try:
        # Check if source file exists
        if not os.path.exists(source_path):
            return {
                "success": False,
                "message": f"Source file not found: {source_path}",
                "filepath": None
            }
        
        # Ensure the directory exists
        os.makedirs(os.path.dirname(os.path.abspath(target_path)), exist_ok=True)
        
        # Convert Word to PDF using docx2pdf
        docx2pdf.convert(source_path, target_path)
        
        logger.info(f"Converted Word to PDF: {source_path} -> {target_path}")
        return {
            "success": True,
            "message": "Successfully converted Word to PDF",
            "filepath": target_path
        }
    except Exception as e:
        logger.error(f"Error converting Word to PDF: {str(e)}")
        return {
            "success": False,
            "message": f"Error converting Word to PDF: {str(e)}",
            "filepath": None
        }

# ---- Resources ----

@server.resource("capabilities://")
def get_capabilities() -> Dict[str, Any]:
    """
    Provide information about this MCP server's capabilities.
    
    Returns:
        Dictionary containing capabilities information
    """
    return {
        "name": "Document Operations",
        "version": "0.1.0",
        "description": "Model Context Protocol server for document operations (Word, Excel, PDF)",
        "document_operations": {
            "word": {
                "create": True,
                "edit": True,
                "convert_from_txt": True,
		"extract_text": True
            },
            "excel": {
                "create": True,
                "edit": True,
                "convert_from_csv": True
            },
            "pdf": {
                "create": True,
                "convert_from_word": True
            }
        }
    }

def main():
    """Main entry point for the server."""
    try:
        # Setup logging directory
        log_dir = Path(__file__).parent.parent / "logs"
        log_dir.mkdir(exist_ok=True)
        
        # Log to file instead of stdout
        startup_logger = logging.getLogger("startup")
        startup_logger.setLevel(logging.INFO)
        
        # Make sure startup logger doesn't also log to console
        startup_logger.propagate = False
        
        # Add file handler for startup logs
        startup_log_file = log_dir / "startup.log"
        file_handler = logging.FileHandler(startup_log_file)
        file_handler.setFormatter(logging.Formatter("%(asctime)s [%(levelname)s] %(message)s"))
        startup_logger.addHandler(file_handler)
        
        # Log startup information to file only
        startup_logger.info("Starting Document Operations MCP Server...")
        startup_logger.info(f"Python version: {sys.version}")
        startup_logger.info(f"Python executable: {sys.executable}")
        startup_logger.info(f"Working directory: {os.getcwd()}")
        startup_logger.info(f"Logs directory: {log_dir}")
        
        # Verify environment
        if sys.prefix == sys.base_prefix:
            startup_logger.warning("Not running in a viratual environment")
            
        startup_logger.info("Server is ready to accept connections from Claude Desktop!")
        
        # Run the server
        server.run()
    except Exception as e:
        logger.error(f"Error starting server: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        sys.exit(1)

if __name__ == "__main__":
    main()
